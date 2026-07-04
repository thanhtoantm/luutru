import streamlit as st
import pandas as pd
import os
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Cấu hình file lưu trữ dữ liệu
DATA_FILE = "du_lieu_gop_y.csv"
WORD_FILE = "tonghop.docx"

def luu_y_kien(mang_bo_phan, ho_ten, chuc_vu, noi_dung):
    new_data = {
        "Thời gian": [datetime.now().strftime("%d/%m/%Y %H:%M:%S")],
        "Mảng/Bộ phận": [mang_bo_phan],
        "Họ và Tên": [ho_ten if ho_ten else "Ẩn danh"],
        "Chức vụ Đảng/Chính quyền": [chuc_vu if chuc_vu else "Đảng viên/Quần chúng"],
        "Nội dung ý kiến": [noi_dung]
    }
    df_new = pd.DataFrame(new_data)
    if os.path.exists(DATA_FILE):
        df_new.to_csv(DATA_FILE, mode='a', header=False, index=False, encoding='utf-8-sig')
    else:
        df_new.to_csv(DATA_FILE, mode='w', header=True, index=False, encoding='utf-8-sig')

def xuat_file_word():
    if not os.path.exists(DATA_FILE):
        return False
    df = pd.read_csv(DATA_FILE, encoding='utf-8-sig')
    doc = Document()
    
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = title_p.add_run("BÁO CÁO TỔNG HỢP Ý KIẾN ĐÓNG GÓP\n")
    run_title.font.name = 'Times New Roman'
    run_title.font.size = Pt(13)
    run_title.bold = True
    
    run_subtitle = title_p.add_run(f"Phục vụ sinh hoạt Chi bộ - Ngày tổng hợp: {datetime.now().strftime('%d/%m/%Y')}\n")
    run_subtitle.font.name = 'Times New Roman'
    run_subtitle.font.size = Pt(12)
    run_subtitle.font.italic = True
    
    doc.add_paragraph("---------------------------\n")

    cac_mang = df["Mảng/Bộ phận"].unique()
    for mang in cac_mang:
        h = doc.add_paragraph()
        run_h = h.add_run(f"I. Ý KIẾN THU THẬP TỪ MẢNG: {mang.upper()}")
        run_h.font.name = 'Times New Roman'
        run_h.font.size = Pt(13)
        run_h.bold = True
        
        df_sub = df[df["Mảng/Bộ phận"] == mang]
        stt = 1
        for index, row in df_sub.iterrows():
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.2)
            run_info = p.add_run(f"{stt}. Đồng chí: {row['Họ và Tên']} ({row['Chức vụ Đảng/Chính quyền']}) - [{row['Thời gian']}]\n")
            run_info.font.name = 'Times New Roman'
            run_info.font.size = Pt(11)
            run_info.bold = True
            
            run_content = p.add_run(f"   Nội dung: {row['Nội dung ý kiến']}\n")
            run_content.font.name = 'Times New Roman'
            run_content.font.size = Pt(11)
            stt += 1
        doc.add_paragraph()
    doc.save(WORD_FILE)
    return True

# --- GIAO DIỆN WEB CẢI TIẾN (KHÔNG DÙNG FORM NHÓM) ---
st.set_page_config(page_title="Lấy ý kiến Chi bộ", page_icon="☭", layout="centered")

st.markdown('<p style="color:red; font-size:17px;"><b>☭ HỆ THỐNG THU THẬP Ý KIẾN </b></p>', unsafe_allow_html=True)
st.markdown('<p style="color:blue; font-size:15px;">Xin chào các đồng chí, vui lòng điền thông tin và ghi nhận sự việc mới.</p>', unsafe_allow_html=True)

# Các ô nhập liệu độc lập
danh_sach_mang = ["Công tác Đảng", "Chuyên môn",  "Tin học", "Sự việc", "Các cuộc họp", "Vấn đề khác"]
mang_selected = st.selectbox("1. Chọn Mảng/Bộ phận đóng góp ý kiến:", danh_sach_mang)

ho_ten = st.text_input("2. Họ và tên (Có thể để trống để giữ bí mật danh tính):", key="input_hoten")
chuc_vu = st.text_input("3. Chức vụ (Đảng/Chính quyền - Ví dụ: Bí thư Chi..):", key="input_chucvu")
noi_dung = st.text_area("4. Nội dung ý kiến đóng góp tâm huyết:", key="input_noidung")

# NÚT GỬI
btn_submit = st.button("🚀 GỬI Ý KIẾN ĐỐNG GÓP", type="primary")

if btn_submit:
    if not noi_dung.strip():
        st.error("Vui lòng nhập nội dung ý kiến, không được bỏ trống!")
    else:
        luu_y_kien(mang_selected, ho_ten, chuc_vu, noi_dung)
        st.success("Gửi ý kiến thành công! Cảm ơn đóng góp của đồng chí.")
        st.balloons() # Hiệu ứng chúc mừng khi gửi thành công

# --- KHÔNG GIAN BẢO MẬT DÀNH CHO CẤP ỦY ---
st.markdown("---")
st.subheader("🔑 Khu vực dành cho Chi ủy (Bảo mật)")

mat_khau_admin = st.text_input("Nhập mật khẩu Admin để tổng hợp dữ liệu:", type="password")

if mat_khau_admin == "Chibo2026":
    st.info("Mật khẩu chính xác. Đồng chí có thể xem dữ liệu và xuất báo cáo Word bên dưới.")
    if os.path.exists(DATA_FILE):
        df_view = pd.read_csv(DATA_FILE, encoding='utf-8-sig')
        st.dataframe(df_view)
        
        if st.button("🔄 Bước 1: Tiến hành tổng hợp vào file tonghop.docx"):
            if xuat_file_word():
                st.success("Đã biên tập và nhóm ý kiến thành công vào file Word!")
                
        if os.path.exists(WORD_FILE):
            with open(WORD_FILE, "rb") as f:
                st.download_button(
                    label="📥 Bước 2: Tải file tonghop.docx về máy tính",
                    data=f,
                    file_name="tonghop.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
    else:
        st.warning("Hiện tại chưa có ý kiến nào được đóng góp.")
elif mat_khau_admin != "":
    st.error("Mật khẩu không đúng! Vui lòng kiểm tra lại.")
